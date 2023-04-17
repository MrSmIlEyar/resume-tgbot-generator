from docx import Document
from docx.shared import Inches

# Создаем документ Word
document = Document()

# Добавляем заголовок
document.add_heading('Резюме', 0)

# Добавляем контактную информацию
name = input("Введите ваше имя: ")
email = input("Введите ваш email: ")
phone = input("Введите ваш номер телефона: ")
address = input("Введите ваш адрес: ")

document.add_paragraph(name, style='Heading 2')
document.add_paragraph(email)
document.add_paragraph(phone)
document.add_paragraph(address)

# Добавляем секцию об образовании
document.add_heading('Образование', level=1)
degree = input("Введите вашу степень образования: ")
university = input("Введите название университета: ")
start_date = input("Введите дату начала обучения: ")
end_date = input("Введите дату окончания обучения: ")

p = document.add_paragraph()
p.add_run(degree + ', ').bold = True
p.add_run(university + ', ').italic = True
p.add_run(start_date + ' - ' + end_date).italic = True

# Добавляем секцию об опыте работы
document.add_heading('Опыт работы', level=1)

company = input("Введите название компании: ")
job_title = input("Введите вашу должность: ")
start_date = input("Введите дату начала работы: ")
end_date = input("Введите дату окончания работы: ")
description = input("Введите описание работы: ")

p = document.add_paragraph()
p.add_run(company + ', ').bold = True
p.add_run(job_title + ', ').italic = True
p.add_run(start_date + ' - ' + end_date).italic = True

p = document.add_paragraph(description)

# Добавляем секцию о навыках
document.add_heading('Навыки', level=1)
skills = input("Введите ваши навыки через запятую: ")
skills_list = skills.split(', ')
skill_table = document.add_table(rows=1, cols=len(skills_list))
hdr_cells = skill_table.rows[0].cells
for i in range(len(skills_list)):
    hdr_cells[i].text = skills_list[i]

# Добавляем фотографию
photo_path = input("Введите путь к фотографии: ")
document.add_picture(photo_path, width=Inches(2), height=Inches(2))

# Сохраняем документ
document.save('resume.docx')