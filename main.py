import logging

import docx
import requests
from aiogram import Bot, Dispatcher, types
from aiogram.utils import executor
from docx import Document
import os
import emoji

import dbConfig
import signup
import markup

# Задаем уровень логов
logging.basicConfig(level=logging.INFO)

# Создаем объект бота
bot = Bot(token="6137850721:AAFSamidQjOex0cjEF5G1SYqtj1knbeZHz8")

# Создаем объект диспетчера
dp = Dispatcher(bot)

# Обработчик команды /start
@dp.message_handler(commands=["start"])
async def cmd_start(message: types.Message):
    signup.signup(message.from_user["username"])
    await message.answer("Привет, я напишу твоё резюме, только лишь расскажи о себе)", reply_markup=markup.MainMenu)

@dp.message_handler(commands=["get_resume"])
async def resume(message: types.Message):
    await message.answer("Здесь будет резюме")

@dp.message_handler()
async def main_dialog(message: types.Message):
    if message.text == "Профиль":
        username = message.from_user["username"]
        supported_user = username.replace('.', '-')
        request = requests.get(dbConfig.dbConfig["databaseURL"] + "/.json" + '?auth=' + dbConfig.Secret)
        data = request.json()
        phone_number = data[supported_user]["Номер телефона"]
        await message.answer(emoji.emojize('👤') + 'Профиль:' + '\n\n' +
                             'Имя пользователя: ' + username + '\n' +
                             'Телефон: ' + phone_number, reply_markup=markup.Profile)

    elif message.text == "Рассказать о себе":
        await message.answer("Тут будет опрос")

    elif message.text == "Получить резюме":
        user_id = message.from_user.id

        user = message.from_user
        user_info = get_user_info(user)
        # сохраняем документ с заданным именем
        filename = f"{user_id}.docx"
        directory = "C:\\Users\\arosl\\PycharmProjects\\resume-tgbot-generator"  # замените на свой путь
        filepath = os.path.join(directory, filename)
        generate_docx_with_array_elements(user_info, filepath)

        # Получаем chat_id пользователя
        chat_id = message.chat.id
        user_id = message.from_user.id
        # Получаем путь к файлу docx
        try:
            docx_path = f'C:\\Users\\arosl\\PycharmProjects\\resume-tgbot-generator\\{user_id}.docx'
            # Открываем файл docx в бинарном режиме
            with open(docx_path, 'rb') as docx:
                # Отправляем документ пользователю
                await bot.send_document(chat_id, docx)
        except:
            await message.answer(f"Вашего файла не обнаружено")

    elif message.text == "Изменить информацию о себе":
        await message.answer("Здесь можно изменить информацию о себе", reply_markup=markup.EditInfo)

    elif message.text == "Изменить фото":
        await message.answer("Здесь можно изменить фото", reply_markup=markup.EditPhoto)

    elif message.text == "Изменить текст":
        await message.answer("Здесь можно изменить текст, который будет в резюме", reply_markup=markup.EditText)

    elif message.text == "Главное меню":
        await message.answer("Главное меню", reply_markup=markup.MainMenu)

    else:
        await message.answer("Мне нужны команды...")


def generate_docx_with_array_elements(array, direct):
    # Создаем объект документа
    document = Document()

    # Добавляем заголовок документа
    document.add_heading('Information about User', level=0)

    # Добавляем элементы массива на новые строки
    for item in array:
        paragraph = document.add_paragraph(str(item))
        paragraph.style = document.styles['List Bullet']

    # Сохраняем документ в файл
    document.save(direct)


def get_user_info(user: types.User):
    user_info = [
        f"User ID: {user.id}",
        f"Username: {user.username}",
        f"First name: {user.first_name}",
        f"Last name: {user.last_name}",
        f"Language: {user.language_code}",
        f"Is bot: {user.is_bot}"
    ]
    return user_info


# Запускаем бота
if __name__ == "__main__":
    executor.start_polling(dp, skip_updates=True)

def creatingadocx(x):
    doc = docx.Document()
    doc.add_heading('Резюме', 0)
    par1 = doc.add_paragraph(x)
    return doc.save('resume.docx')
# f = creatingadocx(x)



